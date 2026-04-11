"""Generate SL dbo Tables Audit as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

def add_status_para(doc, label, status, color=None):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(status)
    r2.font.size = Pt(10)
    if color:
        r2.font.color.rgb = color

# ── Title ───────────────────────────────────────────────────────────────
doc.add_heading('Skill-Lync Warehouse — dbo Tables Complete Audit', level=0)

p = doc.add_paragraph()
p.add_run('Date: ').bold = True
p.add_run('2026-04-10\n')
p.add_run('Database: ').bold = True
p.add_run('Skill-lync Warehouse (Microsoft Fabric)\n')
p.add_run('Audited by: ').bold = True
p.add_run('SL Analytics Agent\n')
p.add_run('Purpose: ').bold = True
p.add_run('Assess readiness of raw dbo tables for a full pipeline rebuild, modeled after Crio\'s clean 4-table architecture.')

# ── Executive Summary ───────────────────────────────────────────────────
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'The Skill-Lync warehouse has 20 dbo tables + 2 views. Of these:'
)
doc.add_paragraph('7 are CORE — needed for the pipeline rebuild', style='List Bullet')
doc.add_paragraph('6 are SUPPORTING — provide enrichment but have quality issues', style='List Bullet')
doc.add_paragraph('7 are DEAD WEIGHT — unused artifacts consuming storage (combined: ~370M+ rows)', style='List Bullet')

doc.add_heading('Top 5 Structural Problems', level=2)
problems = [
    'Activity_extension has 110M rows but 98.4% are orphans — only 1.74M match current activity data',
    'Call and CallEvent are stale since January 2026 — 3 months of calling data missing',
    'LeadCaptureMessage_Parsed is a 340M-row processing artifact serving no purpose',
    'Source metadata (lead_capture_message_metadata) covers only 1.1% of Lead Capture messages',
    'There is no mx_custom decoding — 64 of 65 activity-level custom fields are being lost',
]
for i, prob in enumerate(problems, 1):
    doc.add_paragraph(f'{i}. {prob}', style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Comparison to Crio: ')
r.bold = True
p.add_run('Crio achieves the same analytical capability with 4 clean tables. SL requires 10+ messy tables to get equivalent data.')

# ── Table Classification ────────────────────────────────────────────────
doc.add_heading('Table Classification Overview', level=1)
add_table(doc,
    ['Table', 'Rows', 'Category', 'Status'],
    [
        ['dbo.ActivityBase', '7,547,361', 'CORE', '✅ Clean'],
        ['dbo.ActivityType', '361', 'CORE', '✅ Clean'],
        ['dbo.Activity_extension', '109,902,700', 'SUPPORTING', '🔴 98.4% orphan rows'],
        ['dbo.Leads', '29,472,540', 'CORE', '⚠️ Bloated (back to 2017)'],
        ['dbo.LeadsExtension', '6,467,256', 'SUPPORTING', '⚠️ 22% coverage only'],
        ['dbo.lead_filtered_view', '4,323,478', 'SUPPORTING', '✅ Clean'],
        ['dbo.[User]', '6,012', 'CORE', '⚠️ Dirty names, test accounts'],
        ['dbo.BDATierClassification', '42', 'CORE', '✅ Clean'],
        ['dbo.SkillLyncSalesData', '999', 'CORE', '⚠️ Manual upload'],
        ['dbo.lead_assignment_history', '973,130', 'CORE', '✅ Clean'],
        ['dbo.LeadCaptureMessage', '10,556,756', 'SUPPORTING', '✅ Clean'],
        ['dbo.lead_capture_message_metadata', '116,958', 'SUPPORTING', '🔴 1.1% coverage'],
        ['dbo.Call', '357,264', 'SUPPORTING', '🔴 STALE (Jan 2026)'],
        ['dbo.CallEvent', '332,241', 'DEAD WEIGHT', '🔴 STALE (Jan 2026)'],
        ['dbo.LeadCaptureMessage_Parsed', '340,322,644', 'DEAD WEIGHT', '🔴 Processing artifact'],
        ['dbo.LeadCaptureMessage_Wide', '25,526,377', 'DEAD WEIGHT', '🔴 Processing artifact'],
        ['dbo.SalesData', '1,010', 'DEAD WEIGHT', '🔴 Legacy'],
        ['dbo.SalesDataSumit', '1,212', 'DEAD WEIGHT', '🔴 Legacy'],
        ['dbo.ProspectStageChange', '155,071', 'DEAD WEIGHT', '🔴 Stale, not populated'],
        ['dbo.ProspectOrder', '3,649,038', 'DEAD WEIGHT', '⚠️ Not used (keep for future)'],
    ]
)

# ── CORE Tables Detail ──────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('CORE Tables — Detailed Assessment', level=1)

# ActivityBase
doc.add_heading('1. dbo.ActivityBase — Raw Activity Log', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '7,547,361'],
    ['Unique IDs', '7,547,361 (no duplicates)'],
    ['Date Range', '2025-07-01 → 2026-04-09'],
    ['Status', '✅ Clean'],
    ['Crio Equivalent', 'dbo.Activity'],
])
doc.add_paragraph()
doc.add_paragraph('One row per activity event. Core columns: id, prospect_id, created_at, type_id (FK to ActivityType), source, source_medium, source_campaign, web_url, form_name, activity_note, IP/geo fields.')
doc.add_paragraph('No issues found. Clean, no duplicates, current through yesterday.', style='List Bullet')
doc.add_paragraph('No mx_custom columns here — those live in Activity_extension.', style='List Bullet')

# ActivityType
doc.add_heading('2. dbo.ActivityType — Activity Code Lookup', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '361'],
    ['Status', '✅ Clean'],
])
doc.add_paragraph()
doc.add_paragraph('Maps type_id → activity_code (integer) + display_name. 361 distinct codes, ~25 actively used. Contains known junk codes: 295 (szcfds), 298 (zxc), 433 (hkj), 434 (sdf), 436 (hgvugvhg hj).')

# Leads
doc.add_heading('3. dbo.Leads — Lead Master', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '29,472,540'],
    ['Date Range', '2017-11-05 → 2026-04-09'],
    ['Active Leads (with activity Jul 2025+)', '~68,000 (0.23%)'],
    ['Status', '⚠️ Functional but bloated'],
    ['Crio Equivalent', 'dbo.Leads'],
])
doc.add_paragraph()
doc.add_paragraph('29.5M leads going back to 2017 — only ~68k (0.23%) have any activity in the analysis window.', style='List Bullet')
doc.add_paragraph('All fields are SNAPSHOT (current state) — they change when the lead is updated in CRM.', style='List Bullet')
doc.add_paragraph('Contains internal/test leads (@skill-lync.com, @cybermindworks.com, @criodo.com, @criodo.co.in).', style='List Bullet')

# User
doc.add_heading('4. dbo.[User] — BDA / User Master', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '6,012'],
    ['Status', '⚠️ Dirty but functional'],
])
doc.add_paragraph()
add_table(doc, ['Role', 'Count', 'Level'], [
    ['se (Sales Executive / BDA)', '5,942', 'Bottom'],
    ['dm (Direct Manager)', '12', 'Middle'],
    ['rsm (Regional Sales Manager)', '2', 'Upper'],
    ['ad (Area Director)', '4', 'Top'],
    ['admin', '52', 'System'],
])
doc.add_paragraph()
doc.add_paragraph('Quality issues:', style='List Bullet')
doc.add_paragraph('Leading/trailing spaces in names', style='List Bullet 2')
doc.add_paragraph('Test accounts: @cybermindworks.com emails, IDs like jjbad1/jjbad2', style='List Bullet 2')
doc.add_paragraph('Platform typo: "ohters" should be "others"', style='List Bullet 2')
doc.add_paragraph('No is_active flag — 5,942 SEs but only 42 in BDATierClassification (active roster)', style='List Bullet 2')

# BDATierClassification
doc.add_heading('5. dbo.BDATierClassification — BDA Performance Tier', level=2)
add_table(doc, ['Tier', 'Active', 'Inactive', 'Total'], [
    ['A (top performers)', '11', '0', '11'],
    ['B', '9', '2', '11'],
    ['C', '7', '5', '12'],
    ['New (recent joinees)', '8', '0', '8'],
    ['Total', '35', '7', '42'],
])
doc.add_paragraph()
doc.add_paragraph('Performance-based tiers. Static snapshot (no monthly history unlike Crio).', style='List Bullet')

# SkillLyncSalesData
doc.add_heading('6. dbo.SkillLyncSalesData — Sales / Enrollment Truth', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '999'],
    ['Unique Emails', '821 (178 have multiple purchases)'],
    ['Date Range', '2025-02-01 → 2026-04-06'],
    ['Total Revenue', '₹4.23 crore (~₹42.3M)'],
    ['Status', '⚠️ Manually maintained'],
])
doc.add_paragraph()
doc.add_paragraph('Manually maintained spreadsheet upload — confirmed by user.', style='List Bullet')
doc.add_paragraph('Joined to Leads via LOWER(TRIM(email)) — lossy match (some emails don\'t join).', style='List Bullet')
doc.add_paragraph('Sales Booking Date column is stale (max Dec 2025); Formatted Sale Date is current (Apr 2026).', style='List Bullet')
doc.add_paragraph('No Realized Amount or Realization % — those only exist in legacy SalesData/SalesDataSumit.', style='List Bullet')

# lead_assignment_history
doc.add_heading('7. dbo.lead_assignment_history — BDA Assignment Log', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '973,130'],
    ['Date Range', '2025-07-01 → 2026-04-09'],
    ['Status', '✅ Clean'],
])
doc.add_paragraph()
doc.add_paragraph('Automated assignment events. Contains: selected_user_id (BDA), assignment_type, team_domain (program category), priority_score1, prospect_star_rank.', style='List Bullet')

# ── SUPPORTING Tables ───────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('SUPPORTING Tables — Quality Issues', level=1)

# Activity_extension
doc.add_heading('8. dbo.Activity_extension — Activity Custom Fields', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '109,902,700'],
    ['Rows matching ActivityBase', '1,739,954 (1.6%)'],
    ['Orphan rows', '108,162,746 (98.4%)'],
    ['Columns', '65 mx_custom fields + owner_id + status'],
    ['Status', '🔴 Severely bloated'],
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('CRITICAL PROBLEM: ')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0, 0)
p.add_run('98.4% of rows (108M) are orphans from before Jul 2025. Only 1.74M rows are usable. Every join to this table scans 110M rows to find 1.7M matches.')

doc.add_paragraph()
doc.add_paragraph('Useful data found per activity type:')
add_table(doc, ['Code', 'Activity', 'Matched Rows', 'owner_id', 'mx_custom1', 'mx_custom2'], [
    ['506', 'Kaleyra Outbound Call', '479,250', 'BDA user_id', 'Trunk phone', 'Prospect phone'],
    ['2500', 'DNP', '342,716', '—', 'Attempt count', 'Count'],
    ['777', 'Priority Call', '84,469', '—', 'BDA email', 'BDA user_id'],
    ['393', 'Tech Demo Scheduled', '3,484', 'Demo Engineer', '—', '—'],
    ['395', 'Tech Demo Conducted', '2,732', 'Demo Engineer', '—', '—'],
    ['1000', 'Page Closed', '587,373', '—', '—', '—'],
])

doc.add_paragraph()
doc.add_paragraph('Recommendation: Filter to only rows matching ActivityBase (110M → 1.7M). Decode key mx_custom columns. Create mx_custom mapping document.', style='List Bullet')

# LCM metadata
doc.add_heading('9. dbo.lead_capture_message_metadata — Parsed Source & Program', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '116,958'],
    ['Coverage of LeadCaptureMessage', '1.1% (116K of 10.5M)'],
    ['Status', '🔴 Extremely low coverage'],
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('CRITICAL PROBLEM: ')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0, 0)
p.add_run('98.9% of Lead Capture events have no source data parsed. When lc_source is NULL, the pipeline defaults to "Email" — artificially inflating the Email source bucket.')

# Call
doc.add_heading('10. dbo.Call — Call Detail Records', level=2)
add_table(doc, ['Property', 'Value'], [
    ['Rows', '357,264'],
    ['Date Range', '2025-07-01 → 2026-01-05'],
    ['Status', '🔴 STALE — 3 months behind'],
])
doc.add_paragraph()
add_table(doc, ['call_status', 'rows', '% of total'], [
    ['dnp', '244,628', '68.5%'],
    ['connected', '64,235', '18.0%'],
    ['(null)', '33,975', '9.5%'],
    ['disconnected', '8,925', '2.5%'],
    ['dnpWithinLimit', '5,501', '1.5%'],
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('CRITICAL PROBLEM: ')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0, 0)
p.add_run('Data pipeline stopped syncing in January 2026. No calling data for Feb, Mar, Apr 2026. This is the ONLY source for call connectivity data. Action required: fix the LSQ → Fabric sync.')

# ── DEAD WEIGHT ─────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('DEAD WEIGHT Tables — Can Be Dropped', level=1)

add_table(doc, ['Table', 'Rows', 'Why It\'s Dead Weight'], [
    ['LeadCaptureMessage_Parsed', '340,322,644', 'Exploded JSON artifact. 340M rows of key-value pairs. Not used by any pipeline.'],
    ['LeadCaptureMessage_Wide', '25,526,377', 'Pivoted version of Parsed. Not used.'],
    ['SalesData', '1,010', 'Legacy. User confirmed SkillLyncSalesData is truth.'],
    ['SalesDataSumit', '1,212', 'Legacy. Richer data but not source of truth.'],
    ['ProspectStageChange', '155,071', 'Stale (Jan 2026). User said ignore.'],
    ['CallEvent', '332,241', 'Stale (Jan 2026). Redundant with Call table.'],
    ['ProspectOrder', '3,649,038', 'Not used by pipeline. Keep for future scoring analysis.'],
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Total dead weight: ~370M rows (68% of all warehouse rows). ')
r.bold = True
p.add_run('Dropping these would free significant Fabric capacity.')

# ── Crio vs SL Comparison ──────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Structural Comparison: SL vs Crio', level=1)

doc.add_heading('Crio\'s Clean 4-Table Foundation', level=2)
add_table(doc, ['Crio Table', 'Rows', 'What It Covers'], [
    ['dbo.Activity', '~20M', 'Activities + decoded mx_custom as named columns'],
    ['dbo.Leads', '2.5M', 'Lead master (clean, recent)'],
    ['dbo.Users', '~2K', 'User master + groups'],
    ['dbo.BDA_Mapping', '~200', 'Monthly BDA tier + hierarchy'],
])

doc.add_heading('SL\'s Equivalent Requirement', level=2)
add_table(doc, ['SL Tables Needed', 'Rows', 'Crio Equivalent', 'Gap'], [
    ['ActivityBase + ActivityType + Activity_extension', '7.5M + 361 + 110M', 'dbo.Activity', 'AE is 110M (98% waste). No mx_custom decoding.'],
    ['Leads + LeadsExtension + lead_filtered_view', '29.5M + 6.5M + 4.3M', 'dbo.Leads', '29M rows (90%+ historical). Extension covers 22%.'],
    ['[User] + BDATierClassification', '6K + 42', 'dbo.Users + BDA_Mapping', 'Dirty names, test accounts, no monthly tier.'],
    ['SkillLyncSalesData', '999', '(sale on activity row)', 'Manual upload. Joined by email (lossy).'],
    ['lead_assignment_history', '973K', '(part of Activity)', 'Clean.'],
    ['LeadCaptureMessage + metadata', '10.5M + 117K', '(part of Activity)', 'Metadata covers only 1.1%.'],
    ['Call', '357K', '(part of Activity)', 'STALE since Jan 2026.'],
])

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('SL needs 10+ tables to achieve what Crio does with 4.')
r.bold = True

# ── Action Items ────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Action Items for Discussion', level=1)

doc.add_heading('🔴 Critical — Must Fix Before Rebuild', level=2)
add_table(doc, ['#', 'Issue', 'Impact', 'Owner'], [
    ['1', 'dbo.Call stopped syncing Jan 2026', 'No calling metrics for Feb/Mar/Apr 2026', 'Data Engineering / LSQ team'],
    ['2', 'Activity_extension has 108M orphan rows', 'Pipeline 60x slower than necessary', 'Data Engineering'],
    ['3', 'lead_capture_message_metadata covers 1.1%', '98.9% of Lead Captures default to Email source', 'Data Engineering'],
])

doc.add_heading('🟡 Recommended — During Rebuild', level=2)
add_table(doc, ['#', 'Issue', 'Impact', 'Owner'], [
    ['4', 'Create mx_custom mapping per activity code', 'Unlocks call duration, attempt counts, BDA data', 'Analytics + Product'],
    ['5', 'Drop dead weight tables (340M+ rows)', 'Frees Fabric capacity', 'Data Engineering'],
    ['6', 'Clean dbo.[User] — trim names, flag test accounts', 'Cleaner BDA reporting', 'Data Engineering'],
    ['7', 'Add is_active flag to User or cross-ref BDATier', 'Distinguish current vs former BDAs', 'Analytics'],
])

doc.add_heading('🟢 Can Proceed Now', level=2)
add_table(doc, ['#', 'What\'s Ready', 'Notes'], [
    ['8', 'Core pipeline tables', 'ActivityBase, Leads, User, BDATier, Sales, Assignment — clean enough'],
    ['9', 'Source attribution chain', 'Works but thin (1.1% metadata coverage)'],
    ['10', 'Enrollment join', 'Works but lossy (email matching)'],
])

# ── Storage Summary ─────────────────────────────────────────────────────
doc.add_heading('Storage Summary', level=1)
add_table(doc, ['Category', 'Tables', 'Total Rows', '% of Warehouse'], [
    ['Core', '7', '~38M', '7%'],
    ['Supporting', '6', '~132M', '24%'],
    ['Dead Weight', '7', '~370M', '68%'],
    ['Total', '20', '~540M', '100%'],
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('68% of the warehouse rows are dead weight. ')
r.bold = True
p.add_run('The biggest offender is LeadCaptureMessage_Parsed at 340M rows alone.')

# ── Footer ──────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph('Generated 2026-04-10 by SL Analytics Agent. All row counts and distributions verified via live ODBC queries against the Skill-lync Warehouse.').italic = True

# Save
out_path = os.path.join(os.path.dirname(__file__), 'SL_DBO_Tables_Audit.docx')
doc.save(out_path)
print(f"Saved to {out_path}")
